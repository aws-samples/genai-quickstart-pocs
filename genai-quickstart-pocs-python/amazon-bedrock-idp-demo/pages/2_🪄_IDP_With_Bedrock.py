import streamlit as st
import pandas as pd
from pydantic import BaseModel, create_model, Field
from typing import List, Optional, Dict, Any
import json
import boto3
from io import BytesIO
import base64
import PyPDF2
from docx import Document
from PIL import Image
from typing import Union
import csv
import io
from utils.utils_streamlitApp import *

image_path = "./assets/Banner_1.jpg"
site_icon=Image.open("./assets/IDP_Icon2.png")

# Set up the Streamlit app
st.set_page_config(layout="wide",page_icon=site_icon)
#st.title("📄 Unstructured Document & Image Data Extractor")



# --- File Preview Functions ---
def show_pdf_preview(uploaded_file):
    """Display PDF preview with page navigation"""
    pdf_bytes = uploaded_file.read()
    base64_pdf = base64.b64encode(pdf_bytes).decode('utf-8')
    pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="600" type="application/pdf"></iframe>'
    st.markdown(pdf_display, unsafe_allow_html=True)

def show_image_preview(uploaded_file):
    """Display image preview"""
    st.image(uploaded_file, use_container_width=True)

def show_docx_preview(uploaded_file):
    """Display DOCX as formatted text with basic styling"""
    doc = Document(uploaded_file)
    content = []
    for para in doc.paragraphs:
        if para.text.strip():
            style = "font-size:16px;"  # Default style
            if para.style.name.startswith('Heading'):
                style = "font-size:24px; font-weight:bold;"
            content.append(f'<p style="{style}">{para.text}</p>')
    st.markdown("".join(content), unsafe_allow_html=True)

def show_txt_preview(uploaded_file):
    """Display text file with monospace font"""
    text = uploaded_file.read().decode("utf-8")
    st.markdown(f"```\n{text[:5000]}\n```")  # Show first 5000 chars


def show_csv_preview(uploaded_file):
    """Enhanced CSV Preview with more options"""
    df = pd.read_csv(uploaded_file)

    tab1, tab2 = st.tabs(["📈 Data Preview", "🔍 File Info"])

    with tab1:
        num_rows = st.slider("Select number of rows to display", 1, len(df), 5)
        st.dataframe(df.head(num_rows))
    
    with tab2:
        st.metric("Number of Rows", len(df))
        st.metric("Number of Columns", len(df.columns))
        st.write("Column:", list(df.columns))
        



def process_file_content(uploaded_file):
    """Process file content based on its type and return appropriate content for Bedrock"""
    file_type = uploaded_file.name.split('.')[-1].lower()
    
    try:
        if file_type == 'pdf':
            reader = PyPDF2.PdfReader(uploaded_file)
            text = "\n".join([page.extract_text() for page in reader.pages])
            return {"text": text}
        elif file_type == 'docx':
            doc = Document(uploaded_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return {"text": text}
        elif file_type == 'txt':
            text = uploaded_file.read().decode('utf-8')
            return {"text": text}
        elif file_type in ['png', 'jpeg']:
            # Read image file and convert to base64
            image = Image.open(uploaded_file)
            buffered = BytesIO()
            image.save(buffered, format=image.format)
            image_str = uploaded_file.getvalue()
            return {"image": image_str}
        elif file_type == 'csv':
            # Read and return CSV content
            content = uploaded_file.read().decode('utf-8')
            csv_content = uploaded_file.getvalue().decode('utf-8')
            return {"text": csv_content, "dataframe": pd.read_csv(io.StringIO((csv_content)))}

            
        else:
            st.error(f"Unsupported file type: {file_type}")
            return None
    except Exception as e:
        st.error(f"Error processing file: {e}")
        return None
    

# --- Modified File Processing ---
def process_and_preview_file(uploaded_file):
    """Process file and show appropriate preview"""
    file_type = uploaded_file.name.split('.')[-1].lower()
    
    with st.expander(f"📄 Preview: {uploaded_file.name}", expanded=True):
        if file_type == 'pdf':
            show_pdf_preview(uploaded_file)
        elif file_type in ['png', 'jpg', 'jpeg']:
            show_image_preview(uploaded_file)
        elif file_type == 'docx':
            show_docx_preview(uploaded_file)
        elif file_type == 'txt':
            show_txt_preview(uploaded_file)
        elif file_type == "csv":
            show_csv_preview(uploaded_file)
            #df = pd.read_csv(uploaded_file)
            #st.dataframe(df)
        else:
            st.warning(f"Preview not available for {file_type} files")
    
    # Return processed content (unchanged from your original)
    return process_file_content(uploaded_file)


def generate_pydantic_model(field_definitions: List[Dict]) -> BaseModel:
    """Generate a Pydantic model from user-defined fields."""
    model_fields = {}
    for field in field_definitions:
        field_name = field['name']
        field_type = field['type']
        field_desc = field.get('description', '')

        type_mapping = {
            'string': str,
            'integer': int,
            'float': float,
            'boolean': bool
        }

        python_type = type_mapping.get(field_type, str)

        model_fields[field_name] = (
            Optional[python_type],
            Field(None, description=field_desc)
        )

    DynamicModel = create_model('DynamicModel', **model_fields)
    return DynamicModel

def generate_schema_description(model: BaseModel) -> str:
    """Generate a description of the schema for the LLM."""
    schema = model.model_json_schema()
    description = "Extract data matching this exact schema. Follow these rules:\n"
    description += "- Return ONLY JSON that validates against this schema\n"
    description += "- If a field isn't found, use null\n"
    description += "- Convert values to the specified types\n"
    description += "- Dates should be in YYYY-MM-DD format\n\n"
    description += "SCHEMA DEFINITION:\n"
    description += json.dumps(schema, indent=2)
    return description

def extract_data_with_bedrock(file_content: Dict, schema_description: str, uploaded_file) -> Dict:
    """Use Bedrock to extract data from both text and images"""
    prompt = f"""
    Extract structured data from the provided content (which may be text or an image).
    The content may be a document (like a bank statement, receipt, invoice) or an image of such document.

    {schema_description}
    """
    
    system_prompt = "You are a precise data extraction assistant. Extract ONLY the requested data in valid JSON format that matches the provided schema exactly."
    
    bedrock_runtime = boto3.client(service_name='bedrock-runtime')
    #print("File Content is - {}".format(file_content))
    file_extension = uploaded_file.name.split('.')[-1].lower()
    try:
        # Prepare messages based on content type
        messages = [
            {
                "role": "user",
                "content": [
                    {"text": prompt}
                ]
            }
        ]
        
        # Add either image or text content to the message
        if "image" in file_content:

            messages[0]["content"].append({
                "image": {
                    "format": file_extension,
                    "source": {
                        #"mediaType": "image/jpeg",  # Works for PNG too
                        "bytes": file_content["image"]
                    }
                }
            })
        else:
            messages[0]["content"][0]["text"] += f"\nDocument content:\n{file_content['text'][:20000]}"
        
        #print(messages)
        inf_params = {"maxTokens": 4000, "topP": 0.1, "temperature": 0}
        additionalModelRequestFields = {
            "inferenceConfig": {
                "topK": 40
            }
        }
        response = bedrock_runtime.converse(
            modelId="amazon.nova-lite-v1:0",  # Using Sonnet for better multimodal support

            messages=messages,
            system=[{"text": system_prompt}],
            inferenceConfig=inf_params, ### Commented for Nova
            additionalModelRequestFields=additionalModelRequestFields ## Commented for Nova
        )
        
        output = response['output']['message']['content'][0]['text']
        
        if '```json' in output:
            json_str = output.split('```json')[1].split('```')[0]
        else:
            json_str = output.strip()
        
        return json.loads(json_str)
    except Exception as e:
        st.error(f"Error extracting data with Bedrock: {e}")
        return None
    


def display_extracted_data(extracted_data: Union[Dict, List[Dict]], model: BaseModel):
    """Display extracted data handling all cases including nested lists"""
    try:
        st.header("3. Extracted Data")
        
        # Create tabs
        tab1, tab2 = st.tabs(["JSON View", "Table View"])
        
        with tab1:
            st.json(extracted_data)
        
        with tab2:
            def extract_values(data, field_name):
                """Recursively extract values from nested structures"""
                if isinstance(data, dict):
                    return data.get(field_name, None)
                elif isinstance(data, list):
                    return [extract_values(item, field_name) for item in data]
                return None
            
            # Handle the DynamicModel wrapper
            def unwrap_data(data):
                if isinstance(data, dict) and "DynamicModel" in data:
                    return data["DynamicModel"]
                return data
            
            # Normalize all possible input formats
            if isinstance(extracted_data, list):
                
                # Case 1: List of DynamicModel wrappers
                if all(isinstance(x, dict) and "DynamicModel" in x for x in extracted_data):
                    data_list = [unwrap_data(item) for item in extracted_data]
                # Case 2: List of regular items
                else:
                    data_list = extracted_data
            else:
                # Single item cases
                data_list = [unwrap_data(extracted_data)]
                print(data_list)
            
            # Prepare DataFrame data
            rows = []
            for item in data_list:
                row = {}
                for field_name in model.model_fields.keys():
                    # Handle both direct fields and nested lists
                    value = extract_values(item, field_name)
                    
                    # Format for display
                    if isinstance(value, list):
                        value = ", ".join(str(v) for v in value if v is not None)
                    elif isinstance(value, dict):
                        value = json.dumps(value, indent=2)
                    
                    row[field_name] = value
                rows.append(row)
            
            # Create DataFrame
            field_order = list(model.model_fields.keys())
            df = pd.DataFrame(rows)[field_order] if rows else pd.DataFrame()
            
            # Display with better formatting
            if not df.empty:
                st.dataframe(
                    df.style.set_properties(**{
                        'text-align': 'left',
                        'white-space': 'pre-wrap'
                    }),
                    height=min(400, 35 * (len(df) + 1)),
                    use_container_width=True
                )
            else:
                st.warning("No data to display in table view")
            
            # Validation
            try:
                if isinstance(extracted_data, list):
                    [model(**unwrap_data(item)) for item in extracted_data]
                else:
                    model(**unwrap_data(extracted_data))
                st.success("✅ Data validated against schema")
            except Exception as e:
                st.error(f"Validation error: {e}")
    
    except Exception as e:
        st.error(f"Display error: {e}")

def main():
    display_banner(
        banner_path=image_path,
        caption="Intelligent Document Processing - Powered by Amazon Bedrock"
    )
    contact_sidebar()

    # Step 1: File Upload
    st.subheader("Upload Your Document or Image")
    uploaded_file = st.file_uploader(
        "Choose a file (PDF, DOCX, TXT, PNG, JPEG, GIF,WEBP)",
        type=['pdf', 'docx', 'txt', 'png', 'jpeg','gif','webp'],
        key="doc_uploader"
    )
    
    if uploaded_file:
        with st.spinner("Processing..."):
            file_content = process_and_preview_file(uploaded_file)
            st.session_state.file_content = file_content

        # Display uploaded file preview
        with st.expander("📂 View Uploaded File/content", expanded=False):
            #print("File Content is - {}".format(st.session_state.file_content))
            if "image" in st.session_state.file_content:
                st.image(uploaded_file, caption="Uploaded Image")
            else:
                st.text(st.session_state.file_content['text'][:5000] + ("..." if len(st.session_state.file_content['text']) > 5000 else ""))

            
        # Step 2: Define Fields
        st.subheader("Define Fields to Extract")
            
        if 'fields' not in st.session_state:
            st.session_state.fields = [{'name': '', 'type': 'string', 'description': ''}]
            
        current_fields = st.session_state.fields.copy()
            
        cols = st.columns([3, 2, 4, 1])
        cols[0].write("**Field Name**")
        cols[1].write("**Data Type**")
        cols[2].write("**Description (Optional)**")
        cols[3].write("**Action**")

        fields_to_remove = []
            
        for i, field in enumerate(current_fields):
            cols = st.columns([3, 2, 4, 1])

            new_name = cols[0].text_input(
                    f"Field_{i}_name",
                    value=field['name'],
                    key=f"field_name_{i}",
                    label_visibility="collapsed"
                )

            new_type = cols[1].selectbox(
                    f"Field_{i}_type",
                    options=["string", "integer", "float", "boolean"],
                    index=["string", "integer", "float", "boolean"].index(field['type']),
                    key=f"field_type_{i}",
                    label_visibility="collapsed"
                )

            new_description = cols[2].text_input(
                    f"Field_{i}_description",
                    value=field['description'],
                    key=f"field_description_{i}",
                    label_visibility="collapsed"
                )
            
            if cols[3].button("✕", key=f"remove_{i}"):
                fields_to_remove.append(i)
            else:
                st.session_state.fields[i] = {
                        'name': new_name,
                        'type': new_type,
                        'description': new_description
                    }

        for i in sorted(fields_to_remove, reverse=True):
            if i < len(st.session_state.fields):
                st.session_state.fields.pop(i)
            
        if st.button("➕ Add Field"):
            st.session_state.fields.append({'name': '', 'type': 'string', 'description': ''})

        # Clear all button
        if st.button("🗑️ Clear All", type="secondary"):
            st.session_state.fields = [{'name': '', 'type': 'string', 'description': ''}]
            st.session_state.extracted_data = None
            st.rerun()
            
        if any(field.get('name') for field in st.session_state.fields):
            try:
                model = generate_pydantic_model(st.session_state.fields)
                schema_description = generate_schema_description(model)
                    
                with st.expander("View Generated Schema"):
                    st.code(schema_description, language="json")

                if st.button("Extract Data", type="primary"):
                    with st.spinner("Extracting structured data with Bedrock..."):
                        extracted_data = extract_data_with_bedrock(
                            st.session_state.file_content, 
                            schema_description,
                            uploaded_file
                        )
                        
                    if extracted_data:
                        display_extracted_data(extracted_data,model)

            except Exception as e:
                st.error(f"Error generating schema: {e}")

if __name__ == "__main__":
    main()